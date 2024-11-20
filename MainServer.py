from flask import Flask, render_template_string, request, send_file
import ollama
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
from io import BytesIO

app = Flask(__name__)

# Updated HTML Template
HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>eduai Textbook Generator</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            margin: 0;
            padding: 0;
            background: #202021;
            color: #FFFFFF;
        }
        header {
            text-align: center;
            padding: 1rem 0;
            font-size: 1.5rem;
            background: #202021;
        }
        header span {
            color: #FFFFFF;
        }
        header span.ai {
            color: #0D72E4;
        }
        main {
            display: flex;
            justify-content: center;
            align-items: center;
            height: 100vh;
            margin: 0;
        }
        .content-block {
            background: #333333;
            border-radius: 12px;
            padding: 3rem 2rem;
            width: 90%;
            max-width: 400px;
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.5);
            text-align: center;
        }
        .content-block input,
        .content-block button {
            display: block;
            width: 80%;  /* Same width for both the input field and the button */
            padding: 1rem;
            margin: 1rem auto;
            border: none;
            border-radius: 8px;
            font-size: 1rem;
        }

        .content-block input {
            background: #202021;
            color: #FFFFFF;
            border: 1px solid #FFFFFF;
            margin: 1rem auto;
            display: block;
            width: 80%; /* Adjusted width */
        }
        .content-block input::placeholder {
            color: #CCCCCC;
        }
        .content-block button {
            background: #0D72E4;
            color: #FFFFFF;
            cursor: pointer;
        }
        .content-block button:hover {
            background: #0B5FC1;
        }
        #loading-message {
            display: none;
            font-size: 1.5rem;
            color: #CCCCCC;
        }
        footer {
            text-align: center;
            margin-top: 2rem;
            font-size: 0.8rem;
            color: #777;
        }
    </style>
    <script>
        function showLoading() {
            const form = document.getElementById('textbook-form');
            const loadingMessage = document.getElementById('loading-message');
            
            // Hide the form and show the loading message
            form.style.display = 'none';
            loadingMessage.style.display = 'block';

            // Submit the form
            const formData = new FormData(form);
            fetch('/generate', {
                method: 'POST',
                body: formData,
            })
            .then(response => response.blob())
            .then(blob => {
                const url = window.URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.style.display = 'none';
                a.href = url;
                a.download = formData.get('topic') + '_textbook.docx';
                document.body.appendChild(a);
                a.click();
                window.URL.revokeObjectURL(url);
                loadingMessage.textContent = "Download completed!";
            })
            .catch(() => {
                loadingMessage.textContent = "An error occurred while generating the textbook.";
            });
        }
    </script>
</head>
<body>
    <header>
        <h1><span>edu</span><span class="ai">ai</span> Textbook Generator</h1>
    </header>
    <main>
        <div class="content-block">
            <form id="textbook-form" onsubmit="event.preventDefault(); showLoading();">
                <label for="topic">Enter the Topic for the Textbook:</label>
                <input type="text" id="topic" name="topic" placeholder="E.g., Physics, Biology, History" required>
                <button type="submit">Generate Textbook</button>
            </form>
            <p id="loading-message">Please wait, your textbook is being generated...</p>
        </div>
    </main>
    <footer>
        <p>© 2024 <span>edu</span><span class="ai">ai</span>. All Rights Reserved.</p>
    </footer>
</body>
</html>
"""

# LLM Functionality
def generate_textbook_title(topic):
    prompt = f"Generate a suitable title  only one title and only repond with the one title nothing else not a single word other than a title for a textbook about {topic}. Avoid using symbols or unnecessary text."
    response = ollama.chat(model="llama3.2", messages=[{"role": "user", "content": prompt}])
    return response['message']['content']

def generate_sections_from_ollama(topic):
    prompt = f"Generate an outline for a textbook on {topic}. Only provide section titles and subsections as a clean list."
    response = ollama.chat(model="llama3.2", messages=[{"role": "user", "content": prompt}])
    return response['message']['content']

def generate_section_content_from_ollama(section_title):
    prompt = f"Write detailed content for the section titled: {section_title}. Only include the actual content."
    response = ollama.chat(model="llama3.2", messages=[{"role": "user", "content": prompt}])
    return response['message']['content']

def create_textbook_docx(topic, sections, title):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Cover Page
    doc.add_paragraph().add_run("\n\n\n")
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_paragraph.add_run(title.upper())
    title_run.font.size = Pt(36)
    title_run.bold = True
    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle_paragraph.add_run("Generated by eduai")
    subtitle_run.font.size = Pt(16)
    subtitle_run.italic = True
    doc.add_page_break()

    # Add Sections
    for idx, section in enumerate(sections.split("\n\n"), 1):
        doc.add_heading(f"Section {idx}: {section.strip()}", level=1)
        content = generate_section_content_from_ollama(section.strip())
        doc.add_paragraph(content)
        doc.add_page_break()

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# Flask Routes
@app.route("/", methods=["GET"])
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route("/generate", methods=["POST"])
def generate_textbook():
    topic = request.form.get("topic")
    title = generate_textbook_title(topic)
    sections = generate_sections_from_ollama(topic)
    docx_file = create_textbook_docx(topic, sections, title)

    # Return the DOCX as a download
    return send_file(
        docx_file,
        as_attachment=True,
        download_name=f"{topic}_textbook.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if __name__ == "__main__":
    app.run(debug=True)
